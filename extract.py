#!/usr/bin/env python3
"""
Extract plant catalog data from XLSX and DOCX files.
Creates:
- site/data/plants.json (main catalog)
- site/data/articles/{slug}/article.json (article text)
- site/data/articles/{slug}/img/ (extracted images)
"""

import json
import os
import re
from pathlib import Path
from io import BytesIO

import openpyxl
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image
# Paths
GARDEN_DIR = Path("d:\\garden")
SITE_DIR = GARDEN_DIR / "site"
DATA_DIR = SITE_DIR / "data"
ARTICLES_DIR = DATA_DIR / "articles"
IMG_DIR = ARTICLES_DIR / "img"

XLSX_FILE = GARDEN_DIR / "Каталог растений 12.2025.xlsx"

# Create output directories
DATA_DIR.mkdir(parents=True, exist_ok=True)
ARTICLES_DIR.mkdir(parents=True, exist_ok=True)
IMG_DIR.mkdir(parents=True, exist_ok=True)


# Cyrillic to Latin transliteration mapping
CYRILLIC_TO_LATIN = {
    'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'yo',
    'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
    'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
    'ф': 'f', 'х': 'h', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'sch',
    'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya'
}

def transliterate_ru_to_lat(text):
    """Simple Cyrillic to Latin transliteration."""
    result = []
    for char in text.lower():
        if char in CYRILLIC_TO_LATIN:
            result.append(CYRILLIC_TO_LATIN[char])
        elif ord(char) < 128:  # ASCII character
            result.append(char)
    return ''.join(result)

def slugify(text):
    """Convert text to URL-friendly ASCII slug."""
    if not text:
        return "unnamed"
    # Transliterate Cyrillic to Latin
    slug = transliterate_ru_to_lat(text)
    slug = slug.lower().strip()
    # Remove special characters, keep only alphanumeric and hyphens
    slug = re.sub(r"[^\w\-]", "", slug)
    slug = re.sub(r"\-+", "-", slug)
    return slug.strip("-")


def find_matching_article(name_ru, name_lat):
    """Find matching article for a plant by checking actual article directory structure."""
    if not (ARTICLES_DIR).exists():
        return None

    # Get list of existing articles
    articles = [d.name for d in ARTICLES_DIR.iterdir() if d.is_dir()]
    if not articles:
        return None

    # Try exact match with Russian name slug
    slug_ru = slugify(name_ru)
    if slug_ru in articles:
        return slug_ru

    # Try matching first word of Russian name (handles singular/plural forms)
    first_word = name_ru.split()[0] if name_ru else ""
    if first_word:
        slug_first = slugify(first_word)

        # Direct match
        if slug_first in articles:
            return slug_first

        # Try removing common Russian endings to match singular/plural
        # e.g., "колокольчик" -> "колокольчики"
        for article in articles:
            if article.startswith(slug_first):
                # This article name starts with the same word root
                return article

    # Try matching genus from Latin name
    if name_lat:
        genus = name_lat.split()[0]  # First word is typically genus
        slug_lat = slugify(genus)
        if slug_lat in articles:
            return slug_lat

        # Try partial match with genus
        for article in articles:
            if article.startswith(slug_lat):
                return article

    return None


def parse_xlsx():
    """Parse XLSX catalog and return list of plant dicts."""
    print("Parsing XLSX...")
    wb = openpyxl.load_workbook(XLSX_FILE)
    ws = wb["Перечень"]  # Main inventory sheet

    plants = []

    # Data starts from row 2, columns B through K (indices 1-10)
    for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if not row[1]:  # Skip if column B (name_ru) is empty
            continue

        try:
            name_ru = str(row[1]).strip() if row[1] else ""
            name_lat = str(row[3]).strip() if len(row) > 3 and row[3] else ""

            plant = {
                "id": row_idx - 1,
                "name_ru": name_ru,
                "variety": str(row[2]).strip() if len(row) > 2 and row[2] else "",
                "name_lat": name_lat,
                "family": str(row[4]).strip() if len(row) > 4 and row[4] else "",
                "link": str(row[5]).strip() if len(row) > 5 and row[5] else "",
                "zone": str(row[6]).strip() if len(row) > 6 and row[6] else "",
                "conditions": str(row[7]).strip() if len(row) > 7 and row[7] else "",
                "location": str(row[8]).strip() if len(row) > 8 and row[8] else "",
                "notes": str(row[9]).strip() if len(row) > 9 and row[9] else "",
                "habitat": str(row[10]).strip() if len(row) > 10 and row[10] else "",
            }

            # Find matching article
            article_slug = find_matching_article(name_ru, name_lat)
            plant["slug"] = article_slug if article_slug else slugify(name_ru)
            plant["has_article"] = article_slug is not None
            plant["article_slug"] = article_slug

            plants.append(plant)
        except Exception as e:
            print(f"  Error parsing row {row_idx}: {e}")
            continue

    print(f"  Found {len(plants)} plants")
    return plants


def extract_images_from_docx(docx_path, output_dir):
    """Extract all images from DOCX file."""
    doc = Document(docx_path)
    images = []
    img_num = 1

    # Extract from inline shapes in paragraphs
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_part = rel.target_part
                image_bytes = image_part.blob

                # Determine format
                content_type = image_part.content_type
                ext = "jpg"
                if "png" in content_type:
                    ext = "png"
                elif "gif" in content_type:
                    ext = "gif"

                # Save image (resize for web)
                img = Image.open(BytesIO(image_bytes))
                # Resize to max width 1200px, maintain aspect ratio
                img.thumbnail((1200, 1200), Image.Resampling.LANCZOS)

                img_path = output_dir / f"{img_num:03d}.jpg"
                img.convert("RGB").save(img_path, "JPEG", quality=85)
                images.append(f"{img_num:03d}.jpg")
                img_num += 1
            except Exception as e:
                print(f"    Error extracting image: {e}")

    return images


def classify_paragraph(text, index):
    """Classify paragraph type using heuristics."""
    if not text:
        return None

    # Split by embedded newlines first
    if '\n' in text:
        lines = text.split('\n')
        result = []
        for line in lines:
            line = line.strip()
            if line:
                result.append(classify_paragraph(line, index))
        return result  # Return list of classifications

    text_len = len(text)
    text_lower = text.lower()

    # Check for DOCX style Heading (would be handled by para.style.name in full implementation)
    # For now use heuristics only

    # 1. Caption detection (must be before latin_name)
    if any(word in text_lower for word in ['фотография', 'фото', 'иллюстрация', 'рисунок', 'описание']):
        return {'type': 'caption', 'text': text}

    # 2. Latin name detection: ≤50 chars, mostly Latin letters + spaces/numbers/quotes
    if text_len <= 50:
        latin_count = sum(1 for c in text if c.isalpha() and ord(c) < 128)
        total_alpha = sum(1 for c in text if c.isalpha())
        if total_alpha > 0 and latin_count / total_alpha > 0.7:  # 70% ASCII latin
            return {'type': 'latin_name', 'text': text}

    # 3. Heading detection: ≤60 chars, no ending punctuation, looks like a title
    if text_len <= 60:
        # Check if it looks like a section heading
        has_period = '.' in text
        has_comma = ',' in text
        has_question = '?' in text
        has_digit_in_middle = any(text[i].isdigit() for i in range(1, len(text) - 1) if i < len(text) - 1)

        if not (has_period or has_comma or has_question or has_digit_in_middle):
            # Further check: if it's a known heading keyword
            heading_keywords = [
                'морфология', 'описание', 'характеристика', 'свойства',
                'распространение', 'экология', 'применение', 'выращивание', 'размножение',
                'болезни', 'вредители', 'уход', 'посадка', 'полив', 'удобрение',
                'сорта', 'виды', 'подвиды', 'разновидность', 'классификация',
                'история', 'интересные факты', 'использование', 'галерея',
                'хозяйственное значение', 'исторические сведения', 'ботаническое описание'
            ]
            if any(keyword in text_lower for keyword in heading_keywords):
                return {'type': 'heading', 'text': text}
            # Or if it's short and doesn't look like list item (has capital first letter, no "-" at start)
            elif not text.startswith('-') and text[0].isupper():
                return {'type': 'heading', 'text': text}

    # 4. Default: body text
    return {'type': 'body', 'text': text}


def parse_docx():
    """Parse all DOCX files and extract text + images."""
    print("Parsing DOCX files...")

    docx_files = sorted(GARDEN_DIR.glob("*.docx"))
    articles_data = {}

    for docx_path in docx_files:
        print(f"  {docx_path.name}...")

        try:
            doc = Document(docx_path)
            title = docx_path.stem
            slug = slugify(title)

            # Create article directory
            article_dir = ARTICLES_DIR / slug
            article_img_dir = article_dir / "img"
            article_img_dir.mkdir(parents=True, exist_ok=True)

            # Extract images first
            images = extract_images_from_docx(docx_path, article_img_dir)

            # Extract paragraphs with classification
            paragraphs = []
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue

                # Get style name from DOCX
                style_name = para.style.name if para.style else 'Normal'

                # Classify paragraph
                classified = classify_paragraph(text, idx)

                # Handle list of classifications (from newline splits)
                if isinstance(classified, list):
                    paragraphs.extend(classified)
                elif classified:
                    # Override type for Heading styles
                    if style_name.startswith('Heading'):
                        classified['type'] = 'heading'
                    paragraphs.append(classified)

            # Distribute images throughout paragraphs (insert after captions and between body text)
            paragraphs_with_images = []
            image_idx = 0
            body_para_count = 0

            for para in paragraphs:
                paragraphs_with_images.append(para)

                # Insert image after caption if available
                if para.get('type') == 'caption' and image_idx < len(images):
                    paragraphs_with_images.append({
                        'type': 'image',
                        'src': images[image_idx],
                        'caption': para.get('text', '')
                    })
                    image_idx += 1

                # Or insert image after every 8-12 body paragraphs
                elif para.get('type') == 'body':
                    body_para_count += 1
                    if body_para_count >= 10 and image_idx < len(images):
                        paragraphs_with_images.append({
                            'type': 'image',
                            'src': images[image_idx],
                            'caption': 'Фотография растения'
                        })
                        image_idx += 1
                        body_para_count = 0

            paragraphs = paragraphs_with_images

            # Create article JSON
            article = {
                "title": title,
                "slug": slug,
                "paragraphs": paragraphs,
                "images": images,
                "image_count": len(images),
            }

            # Save article JSON
            article_path = article_dir / "article.json"
            with open(article_path, "w", encoding="utf-8") as f:
                json.dump(article, f, ensure_ascii=False, indent=2)

            articles_data[slug] = article
            print(f"    [OK] {len(paragraphs)} paragraphs, {len(images)} images")

        except Exception as e:
            print(f"    ERROR: {e}")

    return articles_data


def main():
    print("=" * 60)
    print("Plant Catalog Data Extraction")
    print("=" * 60)

    # Parse DOCX first (so articles exist when plants are parsed)
    articles = parse_docx()
    print(f"\n[OK] Extracted {len(articles)} articles with images")

    # Parse XLSX (now articles exist for matching)
    plants = parse_xlsx()

    # Save plants JSON
    plants_path = DATA_DIR / "plants.json"
    with open(plants_path, "w", encoding="utf-8") as f:
        json.dump(plants, f, ensure_ascii=False, indent=2)
    print(f"\n[OK] Saved {len(plants)} plants to {plants_path}")

    # Summary
    plants_with_articles = sum(1 for p in plants if p["has_article"])
    print("\n" + "=" * 60)
    print(f"Summary:")
    print(f"  Total plants:          {len(plants)}")
    print(f"  Plants with articles:  {plants_with_articles}")
    print(f"  Articles extracted:    {len(articles)}")
    print("=" * 60)


if __name__ == "__main__":
    main()
